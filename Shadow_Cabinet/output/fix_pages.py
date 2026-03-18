#!/usr/bin/env python3
"""
Convert Full_Book.md → Final_Book.docx
Professional book formatting for Draft2Digital / print-ready output.

Formatting spec:
  - 6×9 inch trim, 0.75in inside / 0.5in outside / 0.5in top-bottom margins
  - Garamond 12pt body, 1.5 line spacing
  - First paragraph per chapter: no indent; subsequent: 0.5in indent
  - Page breaks before every Part and Chapter
  - Title page, copyright page, TOC
  - Block quotes indented
  - Em-dashes rendered properly
  - Headers (book title odd / author even) and centered page numbers
"""

import re
import os
from docx import Document
from docx.shared import Pt, Inches, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import OxmlElement, parse_xml

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
INPUT_FILE = os.path.join(SCRIPT_DIR, "Full_Book.md")
OUTPUT_FILE = os.path.join(SCRIPT_DIR, "Final_Book.docx")

BOOK_TITLE = "The Shadow Cabinet"
BOOK_SUBTITLE = "Conspiracy, Power, and the Architecture of Hidden Control"
AUTHOR_NAME = "Michael Rodriguez"
COPYRIGHT_YEAR = "2025"

BODY_FONT = "Garamond"
BODY_SIZE = Pt(12)
LINE_SPACING = 1.5
FIRST_LINE_INDENT = Inches(0.5)

PAGE_WIDTH = Inches(6)
PAGE_HEIGHT = Inches(9)
MARGIN_INSIDE = Inches(0.75)
MARGIN_OUTSIDE = Inches(0.5)
MARGIN_TOP = Inches(0.5)
MARGIN_BOTTOM = Inches(0.5)


def parse_markdown(filepath):
    """Parse Full_Book.md into structured blocks."""
    with open(filepath, "r", encoding="utf-8") as f:
        text = f.read()

    # Normalize em-dashes
    text = text.replace(" -- ", " — ")
    text = text.replace("--", "—")

    lines = text.split("\n")
    blocks = []
    i = 0

    # Skip the header lines (title, subtitle, author, separator)
    # Structure: # Title / ## Subtitle / ### Author / --- 
    while i < len(lines):
        line = lines[i].strip()
        if i < 6 and (line.startswith("#") or line == "---" or line == ""):
            i += 1
            continue
        break

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip decorative separators
        if re.match(r"^[-*=]{3,}\s*$", stripped):
            i += 1
            continue

        # Skip "End of Part" lines
        if stripped.startswith("*End of Part"):
            i += 1
            continue

        # Empty line
        if not stripped:
            i += 1
            continue

        # Block quote (lines starting with >)
        if stripped.startswith(">"):
            quote_text = stripped.lstrip("> ").strip()
            i += 1
            while i < len(lines) and lines[i].strip().startswith(">"):
                quote_text += " " + lines[i].strip().lstrip("> ").strip()
                i += 1
            blocks.append(("blockquote", quote_text))
            continue

        # Headings
        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            heading_text = stripped.lstrip("#").strip()
            blocks.append((f"h{level}", heading_text))
            i += 1
            continue

        # Bold summary lines (e.g., **Summary:** ...)
        if stripped.startswith("**") and ":**" in stripped:
            blocks.append(("body", stripped))
            i += 1
            continue

        # Regular paragraph - collect continuation lines
        para_lines = [stripped]
        i += 1
        while i < len(lines):
            next_line = lines[i].strip()
            if not next_line or next_line.startswith("#") or next_line.startswith(">") or re.match(r"^[-*=]{3,}\s*$", next_line):
                break
            para_lines.append(next_line)
            i += 1
        blocks.append(("body", " ".join(para_lines)))

    return blocks


def set_paragraph_font(paragraph, font_name=BODY_FONT, font_size=BODY_SIZE, bold=False, italic=False, color=None):
    """Set font properties on all runs in a paragraph."""
    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = font_size
        run.font.bold = bold
        run.font.italic = italic
        if color:
            run.font.color.rgb = color


def add_run_with_formatting(paragraph, text, font_name=BODY_FONT, font_size=BODY_SIZE, bold=False, italic=False):
    """Add a run with specific formatting."""
    run = paragraph.add_run(text)
    run.font.name = font_name
    run.font.size = font_size
    run.font.bold = bold
    run.font.italic = italic
    return run


def add_rich_text_paragraph(doc, text, font_name=BODY_FONT, font_size=BODY_SIZE, alignment=None, spacing=None, first_line_indent=None):
    """Add a paragraph with inline Markdown formatting (*italic*, **bold**)."""
    para = doc.add_paragraph()
    if alignment is not None:
        para.alignment = alignment
    pf = para.paragraph_format
    if spacing is not None:
        pf.line_spacing = spacing
    if first_line_indent is not None:
        pf.first_line_indent = first_line_indent
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)

    # Parse inline formatting: **bold**, *italic*
    pattern = r'(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*)'
    last_end = 0
    for m in re.finditer(pattern, text):
        # Add text before this match
        if m.start() > last_end:
            add_run_with_formatting(para, text[last_end:m.start()], font_name, font_size)

        if m.group(2):  # ***bold italic***
            add_run_with_formatting(para, m.group(2), font_name, font_size, bold=True, italic=True)
        elif m.group(3):  # **bold**
            add_run_with_formatting(para, m.group(3), font_name, font_size, bold=True)
        elif m.group(4):  # *italic*
            add_run_with_formatting(para, m.group(4), font_name, font_size, italic=True)
        last_end = m.end()

    # Add remaining text
    if last_end < len(text):
        add_run_with_formatting(para, text[last_end:], font_name, font_size)

    return para


def set_page_size_and_margins(section):
    """Set 6x9 page size with book margins."""
    section.page_width = PAGE_WIDTH
    section.page_height = PAGE_HEIGHT
    section.top_margin = MARGIN_TOP
    section.bottom_margin = MARGIN_BOTTOM
    section.left_margin = MARGIN_INSIDE
    section.right_margin = MARGIN_OUTSIDE


def add_page_break(doc):
    """Add a page break paragraph."""
    para = doc.add_paragraph()
    run = para.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._element.append(br)
    return para


def add_header_footer(section, book_title, author_name):
    """Add headers (title/author) and centered page numbers."""
    # Header
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.text = book_title
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in hp.runs:
        run.font.name = BODY_FONT
        run.font.size = Pt(9)
        run.font.italic = True
        run.font.color.rgb = RGBColor(128, 128, 128)

    # Footer with page number
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run()
    run.font.name = BODY_FONT
    run.font.size = Pt(9)
    # Add PAGE field
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    run._element.append(fld_char_begin)

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    run._element.append(instr_text)

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._element.append(fld_char_end)


def build_toc_entries(blocks):
    """Extract TOC entries from parsed blocks."""
    entries = []
    for btype, text in blocks:
        if btype == "h1" and ("PART" in text.upper() or "INTRODUCTION" == text.upper()):
            entries.append((1, text))
        elif btype == "h2" and not text.startswith("Conspiracy"):
            # h2 under a PART/CHAPTER heading = subtitle; standalone h2 = chapter-level
            pass
    
    # Re-parse to get proper structure
    entries = []
    i = 0
    while i < len(blocks):
        btype, text = blocks[i]
        if btype == "h1":
            upper = text.upper()
            if upper == "INTRODUCTION":
                entries.append((1, "Introduction"))
                if i + 1 < len(blocks) and blocks[i+1][0] == "h2":
                    entries.append((2, blocks[i+1][1]))
            elif "PART" in upper:
                entries.append((1, text))
                if i + 1 < len(blocks) and blocks[i+1][0] == "h2":
                    i += 1  # skip subtitle, already captured in part
            elif "CHAPTER" in upper:
                if i + 1 < len(blocks) and blocks[i+1][0] == "h2":
                    entries.append((2, f"{text}: {blocks[i+1][1]}"))
                else:
                    entries.append((2, text))
            elif upper in ("NOTES & SOURCES", "ACKNOWLEDGMENTS", "ABOUT THE AUTHOR"):
                entries.append((1, text))
        i += 1
    return entries


def create_book():
    """Main function to create the formatted DOCX."""
    print(f"Reading {INPUT_FILE}...")
    blocks = parse_markdown(INPUT_FILE)
    print(f"Parsed {len(blocks)} blocks")

    doc = Document()

    # Set default style
    style = doc.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = BODY_SIZE
    style.paragraph_format.line_spacing = LINE_SPACING

    # Configure page size for default section
    section = doc.sections[0]
    set_page_size_and_margins(section)

    # ========== TITLE PAGE ==========
    # Add vertical spacing
    for _ in range(6):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(BOOK_TITLE.upper())
    run.font.name = BODY_FONT
    run.font.size = Pt(28)
    run.font.bold = True
    p.paragraph_format.space_after = Pt(12)

    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(BOOK_SUBTITLE)
    run.font.name = BODY_FONT
    run.font.size = Pt(16)
    run.font.italic = True
    p.paragraph_format.space_after = Pt(36)

    # Author
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(AUTHOR_NAME)
    run.font.name = BODY_FONT
    run.font.size = Pt(18)
    p.paragraph_format.space_after = Pt(0)

    # ========== COPYRIGHT PAGE ==========
    add_page_break(doc)

    for _ in range(4):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)

    copyright_lines = [
        f"© {COPYRIGHT_YEAR} {AUTHOR_NAME}. All rights reserved.",
        "",
        "The moral right of the author has been asserted. No part of this book may be "
        "reproduced, stored in a retrieval system, or transmitted in any form or by any "
        "means electronic, mechanical, photocopying, recording, or otherwise without "
        "prior written permission from the publisher.",
        "",
        "This book contains the author's opinions and analysis of conspiracy theories, "
        "power structures, and hidden influence. It is sold with the understanding that "
        "the author is not engaged in rendering professional legal or political advice.",
        "",
        f"First edition: {COPYRIGHT_YEAR}",
        "",
        "Published by Rodriguez Press",
    ]

    for line in copyright_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if line:
            run = p.add_run(line)
            run.font.name = BODY_FONT
            run.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(2)

    # ========== TABLE OF CONTENTS ==========
    add_page_break(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("CONTENTS")
    run.font.name = BODY_FONT
    run.font.size = Pt(20)
    run.font.bold = True
    p.paragraph_format.space_after = Pt(24)

    toc_entries = build_toc_entries(blocks)
    for level, title in toc_entries:
        p = doc.add_paragraph()
        if level == 1:
            p.paragraph_format.space_before = Pt(12)
            run = p.add_run(title)
            run.font.name = BODY_FONT
            run.font.size = Pt(12)
            run.font.bold = True
        else:
            p.paragraph_format.left_indent = Inches(0.4)
            p.paragraph_format.space_before = Pt(4)
            run = p.add_run(title)
            run.font.name = BODY_FONT
            run.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(2)

    # ========== BODY CONTENT ==========
    is_first_para_in_chapter = True
    current_section_type = None  # 'part', 'chapter', 'intro', 'backmatter'

    i = 0
    while i < len(blocks):
        btype, text = blocks[i]

        if btype == "h1":
            upper = text.upper()

            if upper == "INTRODUCTION":
                add_page_break(doc)
                p = doc.add_paragraph(text.upper(), style="Heading 1")
                set_paragraph_font(p, BODY_FONT, Pt(22), bold=True)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_after = Pt(6)
                current_section_type = "intro"
                is_first_para_in_chapter = True

                # Check for subtitle (h2 following intro h1)
                if i + 1 < len(blocks) and blocks[i+1][0] == "h2":
                    i += 1
                    sub_text = blocks[i][1]
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run(sub_text)
                    run.font.name = BODY_FONT
                    run.font.size = Pt(14)
                    run.font.italic = True
                    p.paragraph_format.space_after = Pt(24)

            elif "PART" in upper:
                add_page_break(doc)
                p = doc.add_paragraph(text.upper(), style="Heading 1")
                set_paragraph_font(p, BODY_FONT, Pt(24), bold=True)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(72)
                p.paragraph_format.space_after = Pt(6)
                current_section_type = "part"

                # Part subtitle
                if i + 1 < len(blocks) and blocks[i+1][0] == "h2":
                    i += 1
                    sub_text = blocks[i][1]
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run(sub_text)
                    run.font.name = BODY_FONT
                    run.font.size = Pt(16)
                    run.font.italic = True
                    p.paragraph_format.space_after = Pt(36)

            elif "CHAPTER" in upper:
                add_page_break(doc)
                p = doc.add_paragraph(text, style="Heading 2")
                set_paragraph_font(p, BODY_FONT, Pt(20), bold=True)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_after = Pt(6)
                current_section_type = "chapter"
                is_first_para_in_chapter = True

                # Chapter subtitle
                if i + 1 < len(blocks) and blocks[i+1][0] == "h2":
                    i += 1
                    sub_text = blocks[i][1]
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run(sub_text)
                    run.font.name = BODY_FONT
                    run.font.size = Pt(14)
                    run.font.italic = True
                    p.paragraph_format.space_after = Pt(24)

            elif upper in ("NOTES & SOURCES", "ACKNOWLEDGMENTS", "ABOUT THE AUTHOR"):
                add_page_break(doc)
                p = doc.add_paragraph(text, style="Heading 1")
                set_paragraph_font(p, BODY_FONT, Pt(20), bold=True)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_after = Pt(18)
                current_section_type = "backmatter"
                is_first_para_in_chapter = True
            else:
                # Generic h1
                add_page_break(doc)
                p = doc.add_paragraph(text, style="Heading 1")
                set_paragraph_font(p, BODY_FONT, Pt(20), bold=True)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_after = Pt(18)
                is_first_para_in_chapter = True

        elif btype == "h2":
            # Standalone h2 (not consumed as subtitle above)
            p = doc.add_paragraph(text, style="Heading 2")
            set_paragraph_font(p, BODY_FONT, Pt(16), bold=True)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            is_first_para_in_chapter = True

        elif btype == "h3":
            p = doc.add_paragraph(text, style="Heading 3")
            set_paragraph_font(p, BODY_FONT, Pt(14), bold=True)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)

        elif btype == "blockquote":
            p = add_rich_text_paragraph(
                doc, text,
                font_name=BODY_FONT, font_size=Pt(11),
                alignment=WD_ALIGN_PARAGRAPH.LEFT,
                spacing=LINE_SPACING,
            )
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.right_indent = Inches(0.5)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            for run in p.runs:
                run.font.italic = True

        elif btype == "body":
            indent = None if is_first_para_in_chapter else FIRST_LINE_INDENT
            p = add_rich_text_paragraph(
                doc, text,
                font_name=BODY_FONT, font_size=BODY_SIZE,
                spacing=LINE_SPACING,
                first_line_indent=indent,
            )
            is_first_para_in_chapter = False

        i += 1

    # ========== HEADERS AND FOOTERS ==========
    for section in doc.sections:
        set_page_size_and_margins(section)
        add_header_footer(section, BOOK_TITLE, AUTHOR_NAME)

    # ========== SAVE ==========
    doc.save(OUTPUT_FILE)
    print(f"\nSaved: {OUTPUT_FILE}")

    # Quick stats
    total_paras = len(doc.paragraphs)
    h1_count = sum(1 for p in doc.paragraphs if p.style.name == "Heading 1")
    h2_count = sum(1 for p in doc.paragraphs if p.style.name == "Heading 2")
    word_count = sum(len(p.text.split()) for p in doc.paragraphs if p.text.strip())
    print(f"Paragraphs: {total_paras}")
    print(f"Parts/Sections (H1): {h1_count}")
    print(f"Chapters (H2): {h2_count}")
    print(f"Word count: {word_count:,}")
    print(f"Font: {BODY_FONT}")
    print(f"Page size: 6×9 inches")
    print("Done — ready for D2D upload.")


if __name__ == "__main__":
    create_book()
