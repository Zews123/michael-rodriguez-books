from docx import Document

doc = Document("Final_Book.docx")

ns = doc.styles["Normal"]
print(f"Default font: {ns.font.name}, size: {ns.font.size}")

for p in doc.paragraphs:
    if "All rights reserved" in p.text:
        for r in p.runs:
            print(f"Copyright: font={r.font.name}, size={r.font.size}, align={p.alignment}")
        break

for p in doc.paragraphs:
    if "investigative journalist" in p.text:
        for r in p.runs[:1]:
            print(f"About author: font={r.font.name}, size={r.font.size}, align={p.alignment}")
        break

for p in doc.paragraphs:
    if "You used Tencent today" in p.text:
        for r in p.runs[:1]:
            print(f"Body text: font={r.font.name}, size={r.font.size}, align={p.alignment}")
        break

last_texts = [p.text for p in doc.paragraphs if p.text.strip()][-3:]
print(f"Last 3 paragraphs: {last_texts}")
