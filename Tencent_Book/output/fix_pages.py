"""
Fix Final_Book.docx:
1. Page break after Title page (after Author line, para index 2)
2. Page break after Copyright page (after Singapore line, para index 9)
3. Page break after About the Author section (after last about paragraph, para index 16)
4. Subtitle font size = 20pt
5. Keep Garamond everywhere, copyright 10pt centered, about centered 12pt
"""
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_page_break_after(paragraph):
    """Add a page break after a paragraph by setting pageBreakBefore on the NEXT paragraph,
    or by adding a run with a page break to the current paragraph."""
    run = paragraph.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._element.append(br)

doc = Document("Final_Book.docx")

# --- STEP 1: Find key paragraphs by index ---
# Title=0, Subtitle=1, Author=2
# COPYRIGHT PAGE starts at 3, ends at 9 (Singapore)
# About the Author heading=11, last about para is the one before next empty/heading

# Identify paragraph indices
title_idx = None
subtitle_idx = None
author_idx = None
copyright_start = None
copyright_end = None
about_heading_idx = None
about_end_idx = None
intro_idx = None

for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    style = p.style.name
    
    if style == "Title":
        title_idx = i
    elif style == "Subtitle":
        subtitle_idx = i
    elif style == "Author":
        author_idx = i
    elif text == "COPYRIGHT PAGE":
        copyright_start = i
    elif "Singapore" in text and copyright_end is None:
        copyright_end = i
    elif text == "About the Author" and style.startswith("Heading"):
        about_heading_idx = i
    elif text == "INTRODUCTION" and style.startswith("Heading"):
        intro_idx = i
        break

# About section ends just before INTRODUCTION heading
# Find the last non-empty paragraph before intro_idx
about_end_idx = intro_idx - 1
while about_end_idx > about_heading_idx and not doc.paragraphs[about_end_idx].text.strip():
    about_end_idx -= 1

print(f"Title: {title_idx}, Subtitle: {subtitle_idx}, Author: {author_idx}")
print(f"Copyright: {copyright_start}-{copyright_end}")
print(f"About: {about_heading_idx}-{about_end_idx}")
print(f"Intro: {intro_idx}")

# --- STEP 2: Apply styling ---
in_copyright = False
in_about = False

for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    style = para.style.name
    
    # Track sections
    if i == copyright_start:
        in_copyright = True
    if i == about_heading_idx:
        in_copyright = False
        in_about = True
    if i == intro_idx:
        in_about = False
    
    # Set Garamond on ALL runs
    for run in para.runs:
        run.font.name = "Garamond"
    
    # Title page elements: centered
    if style == "Title":
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.font.name = "Garamond"
    
    if style == "Subtitle":
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.font.size = Pt(20)
            run.font.name = "Garamond"
    
    if style == "Author":
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.font.name = "Garamond"
    
    # Copyright: 10pt, centered
    if in_copyright:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.font.size = Pt(10)
    
    # About the Author: centered, 12pt
    elif in_about:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.font.size = Pt(12)
    
    # Normal body: 12pt
    elif not style.startswith("Heading") and not style.startswith("Title") and style != "Subtitle" and style != "Author":
        for run in para.runs:
            run.font.size = Pt(12)

# --- STEP 3: Insert page breaks ---
# After Author (para 2) -> separates title page from copyright
add_page_break_after(doc.paragraphs[author_idx])
print(f"Page break added after Author (idx {author_idx})")

# After last copyright paragraph -> separates copyright from About
add_page_break_after(doc.paragraphs[copyright_end])
print(f"Page break added after Copyright end (idx {copyright_end})")

# After last About paragraph -> separates About from Introduction
add_page_break_after(doc.paragraphs[about_end_idx])
print(f"Page break added after About end (idx {about_end_idx})")

# --- STEP 4: Set default style ---
ns = doc.styles["Normal"]
ns.font.name = "Garamond"
ns.font.size = Pt(12)

# Save
doc.save("Final_Book.docx")
print("Done! Page breaks inserted, styling applied.")
