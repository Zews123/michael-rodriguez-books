from docx import Document
import re

doc = Document("Final_Book.docx")

word_count = sum(len(p.text.split()) for p in doc.paragraphs)
h1 = [p.text for p in doc.paragraphs if p.style.name == "Heading 1"]
h2 = [p.text for p in doc.paragraphs if p.style.name == "Heading 2"]
h3 = [p.text for p in doc.paragraphs if p.style.name == "Heading 3"]

artifacts = [p.text for p in doc.paragraphs if re.search(r"^#{1,6} |^---", p.text)]
copyright_present = any("©" in p.text or "COPYRIGHT PAGE" in p.text for p in doc.paragraphs)

print(f"Word count   : {word_count:,}")
print(f"Parts (H1)   : {len(h1)}")
print(f"Chapters (H2): {len(h2)}")
print(f"Sections (H3): {len(h3)}")
print()

if word_count >= 30000:
    print("OK Word count")
else:
    print(f"FAIL: Word count {word_count} < 30,000")

if h2:
    print("OK Headings detected")
else:
    print("FAIL: No headings")

if not artifacts:
    print("OK No Markdown artifacts")
else:
    print(f"FAIL Artifacts found: {len(artifacts)}")

if copyright_present:
    print("OK Copyright page detected")
else:
    print("WARN No copyright page found")

print()
print("Parts:")
for h in h1:
    print(f"  [{h}]")
print("Chapters (first 10):")
for h in h2[:10]:
    print(f"  [{h}]")
