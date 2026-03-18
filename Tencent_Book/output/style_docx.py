"""
Post-process Final_Book.docx:
1. Set Garamond as the default font (all body text, 12pt)
2. Copyright page paragraphs: Garamond 10pt, centered
3. About the Author heading + text: centered
4. Headings: Garamond
"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document("Final_Book.docx")

# --- Identify copyright and about-author sections ---
# Copyright block: from "COPYRIGHT PAGE" until "About the Author"
# About the Author block: from "About the Author" heading until "INTRODUCTION"

in_copyright = False
in_about = False

for para in doc.paragraphs:
    text = para.text.strip()

    # Detect sections
    if text == "COPYRIGHT PAGE":
        in_copyright = True
    if text == "About the Author" and para.style.name.startswith("Heading"):
        in_copyright = False
        in_about = True
    if text == "INTRODUCTION" and para.style.name.startswith("Heading"):
        in_about = False

    # Apply font to ALL paragraphs
    for run in para.runs:
        run.font.name = "Garamond"

    # Copyright section: 10pt + centered
    if in_copyright:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.font.size = Pt(10)

    # About the Author section: centered, normal size (12pt)
    elif in_about:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.font.size = Pt(12)

    # Normal text: 12pt
    elif not para.style.name.startswith("Heading") and not para.style.name.startswith("Title"):
        for run in para.runs:
            run.font.size = Pt(12)

# Also set Garamond on heading runs (keep their existing sizes)
for para in doc.paragraphs:
    if para.style.name.startswith("Heading") or para.style.name.startswith("Title"):
        for run in para.runs:
            run.font.name = "Garamond"

# Set default font in the document's style
style = doc.styles["Normal"]
style.font.name = "Garamond"
style.font.size = Pt(12)

doc.save("Final_Book.docx")
print("Done. Garamond applied, copyright centered at 10pt, About the Author centered at 12pt.")
