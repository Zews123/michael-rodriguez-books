"""
Post-process Final_Book.docx for Draft2Digital:
1. Set Garamond as default font (12pt body text)
2. Copyright page: Garamond 10pt, centered
3. About the Author: Garamond 12pt, centered
4. Subtitle: 20pt
5. Page breaks between: Title -> Copyright -> About -> Introduction
"""
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def add_page_break_after(paragraph):
    run = paragraph.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._element.append(br)


doc = Document("Final_Book.docx")

title_idx = None
subtitle_idx = None
author_idx = None
copyright_start = None
copyright_end = None
about_heading_idx = None
about_end_idx = None
intro_idx = None

for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    style = p.style.name

    if style == "Title":
        title_idx = i
    elif style == "Subtitle":
        subtitle_idx = i
    elif style == "Author":
        author_idx = i
    elif text == "COPYRIGHT PAGE" and copyright_start is None:
        copyright_start = i
    elif text == "About the Author" and style.startswith("Heading"):
        about_heading_idx = i
        copyright_end = i - 1
        while copyright_end > (copyright_start or 0) and not doc.paragraphs[copyright_end].text.strip():
            copyright_end -= 1
    elif text == "INTRODUCTION" and style.startswith("Heading"):
        intro_idx = i
        about_end_idx = i - 1
        while about_end_idx > (about_heading_idx or 0) and not doc.paragraphs[about_end_idx].text.strip():
            about_end_idx -= 1
        break

print(f"Title: {title_idx}, Subtitle: {subtitle_idx}, Author: {author_idx}")
print(f"Copyright: {copyright_start}-{copyright_end}")
print(f"About: {about_heading_idx}-{about_end_idx}")
print(f"Intro: {intro_idx}")

in_copyright = False
in_about = False

for i, para in enumerate(doc.paragraphs):
    style = para.style.name

    if i == copyright_start:
        in_copyright = True
    if i == about_heading_idx:
        in_copyright = False
        in_about = True
    if i == intro_idx:
        in_about = False

    for run in para.runs:
        run.font.name = "Garamond"

    if style in ("Title", "Author"):
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if style == "Subtitle":
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.font.size = Pt(20)

    elif in_copyright:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.font.size = Pt(10)

    elif in_about:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.font.size = Pt(12)

    elif not style.startswith("Heading") and style not in ("Title", "Subtitle", "Author"):
        for run in para.runs:
            run.font.size = Pt(12)

if author_idx is not None:
    add_page_break_after(doc.paragraphs[author_idx])
    print(f"Page break after Title page (idx {author_idx})")

if copyright_end is not None:
    add_page_break_after(doc.paragraphs[copyright_end])
    print(f"Page break after Copyright (idx {copyright_end})")

if about_end_idx is not None:
    add_page_break_after(doc.paragraphs[about_end_idx])
    print(f"Page break after About the Author (idx {about_end_idx})")

ns = doc.styles["Normal"]
ns.font.name = "Garamond"
ns.font.size = Pt(12)

doc.save("Final_Book.docx")
print("Done. Professional styling applied.")
